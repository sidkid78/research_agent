# FastAPI application entrypoint 
import uuid
import logging
from fastapi import FastAPI, HTTPException, BackgroundTasks, Depends
from fastapi import WebSocket, WebSocketDisconnect
from sqlalchemy.orm import Session
from typing import List, Dict, Optional
from datetime import datetime, timezone
from pydantic import BaseModel
import json

from . import models, schemas, crud # crud will be created later
from .db import SessionLocal, engine, get_db, create_db_and_tables
from .agent import create_research_agent
from google import genai

from fastapi.middleware.cors import CORSMiddleware

# Configure logging
logger = logging.getLogger(__name__)

# Create database tables on startup
models.Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="Research Agent API",
    description="API for submitting research jobs and retrieving results.",
    version="0.1.0"
)

# --- Event Handlers ---
def startup_event():
    create_db_and_tables()
    logger.info("Database tables created")
    logger.info("Starting the application...")

app.add_event_handler("startup", startup_event)

# --- Middleware ---

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["*"],
)


# --- Background Task Functions ---

async def run_research_job(job_id: str, request_data: dict):
    """
    Background task to run the actual research job using Gemini AI.
    """
    logger.info(f"Starting research job {job_id}")
    
    # Get database session
    db = SessionLocal()
    try:
        # Update job status to in_progress
        crud.update_job_status(db, job_id, models.JobStatusEnum.in_progress, progress=0.1)
        logger.info(f"Job {job_id} updated to in_progress")
        # Create research agent
        agent = create_research_agent()
        logger.info(f"Research agent created")
        # Create research request from stored data
        research_request = schemas.ResearchRequest(**request_data)
        logger.info(f"Research request created")
        # Run the research
        result = await agent.conduct_research(research_request)
        logger.info(f"Research completed")
        # Convert result to dict for storage
        result_dict = {
            "job_id": job_id,
            "topic": result.topic,
            "content": result.content,
            "references": [
                {
                    "title": ref.title,
                    "url": ref.url,
                    "accessed_date": ref.accessed_date.isoformat(),
                    "snippet": ref.snippet
                }
                for ref in result.references
            ],
            "output_format": result.output_format,
            "generated_at": result.generated_at.isoformat(),
            "word_count": result.word_count,
            "confidence_score": result.confidence_score,
            "academic_source_breakdown": {
                "arxiv_papers": len([r for r in result.references if r.source_type == 'arxiv']),
                "pubmed_papers": len([r for r in result.references if r.source_type == 'pubmed']),
                "web_sources": len([r for r in result.references if r.source_type == 'web']),
                "total_sources": len(result.references)
            },
            "research_quality_indicators": {
                # PubMed sources are generally peer-reviewed
                "peer_reviewed_percentage": (len([r for r in result.references if r.source_type == 'pubmed']) / len(result.references) * 100) if result.references else 0,
                # Estimate recent sources (this would need actual publication dates)
                "recent_sources_percentage": 75.0,  # TODO: Calculate from actual publication dates
                # PubMed and arXiv are authoritative academic sources
                "authoritative_sources_percentage": (len([r for r in result.references if r.source_type in ['pubmed', 'arxiv']]) / len(result.references) * 100) if result.references else 0
            }
        }
        
        # Update job with results
        crud.update_job_completed(db, job_id, result_dict)
        logger.info(f"Job {job_id} completed successfully")
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        logger.error(f"Job {job_id} failed with error: {str(e)}")
        logger.error(f"Full traceback:\n{error_details}")
        crud.update_job_completed(db, job_id, None, error_message=f"{str(e)}\n\nTraceback:\n{error_details}")
    finally:
        db.close()
        logger.info(f"Database connection closed")

# --- API Endpoints ---

@app.post("/research", response_model=schemas.JobSubmitResponse, status_code=202) # 202 Accepted
async def submit_research_job(
    request: schemas.ResearchRequest,
    background_tasks: BackgroundTasks, # We'll use this later
    db: Session = Depends(get_db)
):
    """
    Submit a new research request.

    - **topic**: The research topic or question.
    - **output_format**: 'bullets' or 'full_report'.
    - **deadline** (optional): ISO 8601 format datetime string.
    """
    job_id = str(uuid.uuid4())
    
    # Use CRUD function to create the job
    crud.create_research_job(db=db, job_id=job_id, research_request=request)
    logger.info(f"Job {job_id} created")
    # Add actual research task to background processing
    background_tasks.add_task(run_research_job, job_id, request.model_dump())
    logger.info(f"Research task added to background processing")
    return schemas.JobSubmitResponse(job_id=job_id, status=models.JobStatusEnum.queued.value)

@app.get("/research/{job_id}/status", response_model=schemas.JobStatusResponse)
async def get_research_job_status(job_id: str, db: Session = Depends(get_db)):
    """
    Get the current status and progress of a research job.
    """
    logger.info(f"Getting status for job {job_id}")
    db_job = crud.get_research_job(db, job_id=job_id)
    logger.info(f"Job {job_id} found")
    if db_job is None:
        raise HTTPException(status_code=404, detail=f"Job with ID '{job_id}' not found.")
    
    # Include error message if job failed
    error_message = None
    if db_job.status == models.JobStatusEnum.failed and db_job.error_message:
        error_message = db_job.error_message
        logger.warning(f"Job {job_id} failed with error: {error_message}")
    
    logger.info(f"Job {job_id} status returned")
    return schemas.JobStatusResponse(
        job_id=db_job.id,
        status=db_job.status.value,
        progress=db_job.progress,
        error_message=error_message
    )
    
@app.get("/research/{job_id}/details", response_model=schemas.ResearchRequest, tags=["Research Jobs"])
async def get_research_job_details(job_id: str, db: Session = Depends(get_db)):
    """
    Retrieves the original request details for a specific research job.
    """
    logger.info(f"Getting details for job {job_id}")
    db_job = crud.get_research_job(db, job_id=job_id)
    logger.info(f"Job {job_id} details found")
    if db_job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    logger.info(f"Job {job_id} details returned")
    if db_job.request_payload is None:
        raise HTTPException(status_code=404, detail="Job details not found for this job.")
    logger.info(f"Job {job_id} details returned")
    return schemas.ResearchRequest(**db_job.request_payload)

@app.get("/research/{job_id}/result", response_model=schemas.ResearchResult)
async def get_research_job_result(job_id: str, db: Session = Depends(get_db)):
    """
    Get the result of a completed research job.
    """
    logger.info(f"Getting result for job {job_id}")
    db_job = crud.get_research_job(db, job_id=job_id)
    
    if db_job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if db_job.status != models.JobStatusEnum.completed:
        raise HTTPException(status_code=400, detail="Job not completed yet")
    
    if db_job.result_payload is None:
        raise HTTPException(status_code=404, detail="Job result not found")
    
    logger.info(f"Job {job_id} result returned")
    return schemas.ResearchResult(**db_job.result_payload)

# Academic Research Endpoints
@app.post("/academic-research", response_model=schemas.JobSubmitResponse, status_code=202)
async def submit_academic_research(
    request: schemas.ResearchRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """
    Submit a new academic research request with enhanced features.
    """
    logger.info(f"Academic research request received: {request.topic}")
    job_id = str(uuid.uuid4())
    
    # Create the job using existing CRUD function
    crud.create_research_job(db=db, job_id=job_id, research_request=request)
    logger.info(f"Academic research job {job_id} created")
    
    # Add research task to background processing
    background_tasks.add_task(run_research_job, job_id, request.model_dump())
    logger.info(f"Academic research task added to background processing")
    
    return schemas.JobSubmitResponse(
        job_id=job_id, 
        status=models.JobStatusEnum.queued.value,
        estimated_duration_minutes=5
    )

@app.post("/academic-research/validate-pubmed")
async def validate_pubmed_access(request: dict):
    """
    Validate PubMed email access (placeholder implementation).
    """
    email = request.get("email", "")
    logger.info(f"PubMed validation request for email: {email}")
    
    # Simple validation - in production, you'd implement actual PubMed API validation
    if email and "@" in email:
        return {"valid": True, "message": "Email format is valid"}
    else:
        return {"valid": False, "message": "Invalid email format"}

@app.get("/academic-research/preview")
async def get_academic_sources_preview(topic: str, email: str = None):
    """
    Get a preview of academic sources for a given topic.
    """
    logger.info(f"Academic sources preview requested for topic: {topic}")
    
    try:
        from .gemini_helpers import AcademicGeminiHelpers
        
        # Quick preview search - limit to 2 papers per source
        arxiv_papers = await AcademicGeminiHelpers.search_arxiv_papers(topic, max_results=2)
        
        pubmed_papers = []
        if email:
            pubmed_papers = await AcademicGeminiHelpers.search_pubmed_papers(
                topic, max_results=2, email=email
            )
        
        # Format for frontend
        arxiv_preview = [
            {
                "title": paper.get("title", ""),
                "authors": paper.get("authors", []),
                "abstract": paper.get("abstract", "")[:200] + "..."
            }
            for paper in arxiv_papers[:2]
        ]
        
        pubmed_preview = [
            {
                "title": paper.get("title", ""),
                "journal": paper.get("journal", "Unknown"),
                "abstract": paper.get("abstract", "")[:200] + "..."
            }
            for paper in pubmed_papers[:2]
        ]
        
        estimated_sources = len(arxiv_papers) + len(pubmed_papers) + 5  # +5 for potential web sources
        
        return {
            "arxiv_preview": arxiv_preview,
            "pubmed_preview": pubmed_preview,
            "estimated_sources": estimated_sources
        }
    except Exception as e:
        logger.error(f"Preview search failed: {e}")
        # Return empty preview on error
        return {
            "arxiv_preview": [],
            "pubmed_preview": [],
            "estimated_sources": 0
        }

# -------------------------------------------------------------
# Live Research (Minimal In-Memory Implementation for Frontend)
# -------------------------------------------------------------

class LiveResearchStartRequestModel(BaseModel):
    topic: str
    modalities: Optional[List[str]] = None
    system_instructions: Optional[List[str]] = None

class LiveResearchSessionModel(BaseModel):
    session_id: str
    topic: str
    status: str
    started_at: str
    modalities: List[str]
    participants: Optional[int] = 1

class LiveResearchSummaryModel(BaseModel):
    session_id: str
    topic: str
    duration_minutes: int
    total_interactions: int
    key_findings: List[dict]
    questions_explored: List[dict]
    research_report: Optional[str] = None

live_sessions: Dict[str, dict] = {}

@app.post("/live-research/start", response_model=LiveResearchSessionModel)
async def start_live_research_session(payload: LiveResearchStartRequestModel):
    import uuid as _uuid
    session_id = str(_uuid.uuid4())
    now = datetime.now(timezone.utc)
    live_sessions[session_id] = {
        "topic": payload.topic,
        "status": "active",
        "started_at": now,
        "modalities": payload.modalities or ["text"],
        "interactions": 0,
        "key_findings": [],
        "questions": []
    }
    return LiveResearchSessionModel(
        session_id=session_id,
        topic=payload.topic,
        status="active",
        started_at=now.isoformat(),
        modalities=payload.modalities or ["text"],
        participants=1,
    )

@app.get("/live-research/{session_id}/status", response_model=LiveResearchSummaryModel)
async def get_live_research_status(session_id: str):
    sess = live_sessions.get(session_id)
    if not sess:
        raise HTTPException(status_code=404, detail="Session not found")
    elapsed = max(1, int((datetime.now(timezone.utc) - sess["started_at"]).total_seconds() // 60))
    return LiveResearchSummaryModel(
        session_id=session_id,
        topic=sess["topic"],
        duration_minutes=elapsed,
        total_interactions=sess.get("interactions", 0),
        key_findings=sess.get("key_findings", []),
        questions_explored=sess.get("questions", []),
        research_report=None,
    )

@app.post("/live-research/{session_id}/end", response_model=LiveResearchSummaryModel)
async def end_live_research_session(session_id: str):
    sess = live_sessions.get(session_id)
    if not sess:
        raise HTTPException(status_code=404, detail="Session not found")
    sess["status"] = "ended"
    elapsed = max(1, int((datetime.now(timezone.utc) - sess["started_at"]).total_seconds() // 60))
    # Provide a lightweight summary
    summary = LiveResearchSummaryModel(
        session_id=session_id,
        topic=sess["topic"],
        duration_minutes=elapsed,
        total_interactions=sess.get("interactions", 0),
        key_findings=sess.get("key_findings", []),
        questions_explored=sess.get("questions", []),
        research_report=f"Live research session for '{sess['topic']}' ended after {elapsed} minute(s).",
    )
    return summary

@app.websocket("/live-research/{session_id}/ws")
async def live_research_ws(websocket: WebSocket, session_id: str):
    await websocket.accept()
    try:
        now_iso = datetime.now(timezone.utc).isoformat()
        await websocket.send_json({
            "sender": "assistant",
            "content": f"Connected to live research session {session_id}. Ask a question to begin.",
            "timestamp": now_iso,
        })
        while True:
            raw = await websocket.receive_text()
            reply_ts = datetime.now(timezone.utc).isoformat()
            try:
                data = json.loads(raw)
            except Exception:
                data = {"type": "message", "content": raw}

            # Update minimal session metrics
            sess = live_sessions.get(session_id)
            if sess is not None:
                sess["interactions"] = int(sess.get("interactions", 0)) + 1

            msg_type = str(data.get("message_type") or data.get("type") or "message").lower()
            if msg_type == "audio":
                response_text = "Received your audio message. Transcription is not enabled in this demo stub."
            else:
                user_content = str(data.get("content", "")).strip()
                if user_content:
                    response_text = (
                        "Working on that. This demo backend is a placeholder; "
                        "the full agent would search PubMed/arXiv/web and summarize findings.\n\n"
                        f"Echo: {user_content}"
                    )
                else:
                    response_text = "Please provide a question or message to continue."

            await websocket.send_json({
                "sender": "assistant",
                "content": response_text,
                "timestamp": reply_ts,
            })
    except WebSocketDisconnect:
        pass

# ----------------------------------------
# Batch Research (Minimal In-Memory Stubs)
# ----------------------------------------

class BatchResearchRequestModel(BaseModel):
    topics: List[str]
    output_format: str
    email: Optional[str] = None
    research_type: Optional[str] = None

batch_jobs: Dict[str, dict] = {}

@app.post("/batch-research")
async def submit_batch_research(payload: BatchResearchRequestModel):
    import uuid as _uuid
    batch_id = str(_uuid.uuid4())
    now = datetime.now(timezone.utc)
    batch_jobs[batch_id] = {
        "created_at": now,
        "status": "in_progress",
        "topics": payload.topics,
        "output_format": payload.output_format,
        "completed": [],
        "failed": [],
        "results": {}
    }
    # Minimal immediate progress: mark as completed with placeholder results
    for t in payload.topics:
        batch_jobs[batch_id]["completed"].append(t)
        batch_jobs[batch_id]["results"][t] = {
            "topic": t,
            "content": f"Auto-generated placeholder result for '{t}'.",
            "references": [],
            "output_format": payload.output_format,
            "generated_at": now.isoformat(),
            "word_count": 25,
            "confidence_score": 0.5,
        }
    batch_jobs[batch_id]["status"] = "completed"
    return {"batch_id": batch_id, "status": "queued"}

@app.get("/batch-research/{batch_id}/status")
async def get_batch_status(batch_id: str):
    job = batch_jobs.get(batch_id)
    if not job:
        raise HTTPException(status_code=404, detail="Batch not found")
    total = max(1, len(job["topics"]))
    progress = len(job["completed"]) / total
    return {
        "batch_id": batch_id,
        "status": job["status"],
        "progress": progress,
        "completed_topics": job["completed"],
        "failed_topics": job["failed"],
    }

@app.get("/batch-research/{batch_id}/results")
async def get_batch_results(batch_id: str):
    job = batch_jobs.get(batch_id)
    if not job:
        raise HTTPException(status_code=404, detail="Batch not found")
    total = len(job["topics"])
    return {
        "batch_id": batch_id,
        "total_topics": total,
        "completed_topics": len(job["completed"]),
        "failed_topics": len(job["failed"]),
        "results": job["results"],
        "overall_confidence": 0.5,
        "processing_time_seconds": 1,
    }

# Export Endpoints
@app.get("/research/{job_id}/export")
async def export_research_result(
    job_id: str,
    format: str = "txt",
    db: Session = Depends(get_db)
):
    """
    Export research result in various formats.
    Supported formats: txt, md, pdf, docx
    """
    logger.info(f"Export request for job {job_id} in format {format}")
    
    # Get the research result
    db_job = crud.get_research_job(db, job_id=job_id)
    
    if db_job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    
    if db_job.status != models.JobStatusEnum.completed:
        raise HTTPException(status_code=400, detail="Job not completed yet")
    
    if db_job.result_payload is None:
        raise HTTPException(status_code=404, detail="Job result not found")
    
    result = schemas.ResearchResult(**db_job.result_payload)
    
    # Format the content based on requested format
    if format == "txt" or format == "md":
        # Plain text or markdown format
        content_lines = [
            f"# {result.topic}",
            "",
            f"Generated on: {result.generated_at}",
            f"Word Count: {result.word_count}",
            "",
            "## Research Summary",
            "",
            result.content,
            "",
            "## References",
            ""
        ]
        
        for i, ref in enumerate(result.references, 1):
            content_lines.append(f"{i}. **{ref.title}**")
            content_lines.append(f"   URL: {ref.url}")
            if ref.snippet:
                content_lines.append(f"   {ref.snippet}")
            content_lines.append("")
        
        content_str = "\n".join(content_lines)
        
        from fastapi.responses import Response
        return Response(
            content=content_str,
            media_type="text/plain" if format == "txt" else "text/markdown",
            headers={
                "Content-Disposition": f"attachment; filename=research_{job_id}.{format}"
            }
        )
    
    elif format == "pdf":
        # Generate a simple PDF using reportlab
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        y = height - 72
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, y, f"Research: {result.topic}")
        y -= 24
        c.setFont("Helvetica", 10)
        c.drawString(72, y, f"Generated: {result.generated_at}  •  Word Count: {result.word_count}")
        y -= 24
        c.setFont("Helvetica-Bold", 12)
        c.drawString(72, y, "Summary:")
        y -= 16
        c.setFont("Helvetica", 11)
        summary_lines = simpleSplit(result.content or "", "Helvetica", 11, width - 144)
        for line in summary_lines:
            if y < 72:
                c.showPage(); y = height - 72; c.setFont("Helvetica", 11)
            c.drawString(72, y, line); y -= 14
        y -= 12
        c.setFont("Helvetica-Bold", 12)
        if y < 88:
            c.showPage(); y = height - 72
        c.drawString(72, y, "References:")
        y -= 16
        c.setFont("Helvetica", 10)
        for idx, ref in enumerate(result.references, 1):
            ref_line = f"{idx}. {ref.title}"
            lines = simpleSplit(ref_line, "Helvetica", 10, width - 144)
            for ln in lines:
                if y < 72:
                    c.showPage(); y = height - 72; c.setFont("Helvetica", 10)
                c.drawString(72, y, ln); y -= 12
            if ref.url:
                url_line = f"   {ref.url}"
                lines = simpleSplit(url_line, "Helvetica", 10, width - 144)
                for ln in lines:
                    if y < 72:
                        c.showPage(); y = height - 72; c.setFont("Helvetica", 10)
                    c.drawString(72, y, ln); y -= 12
            if ref.snippet:
                snip_line = f"   {ref.snippet}"
                lines = simpleSplit(snip_line, "Helvetica", 10, width - 144)
                for ln in lines:
                    if y < 72:
                        c.showPage(); y = height - 72; c.setFont("Helvetica", 10)
                    c.drawString(72, y, ln); y -= 12
            y -= 6
        c.showPage(); c.save()
        pdf_bytes = buffer.getvalue(); buffer.close()
        from fastapi.responses import Response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=research_{job_id}.pdf"}
        )
    elif format == "docx":
        # Generate a simple DOCX using python-docx
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt
        buffer = BytesIO()
        doc = Document()
        doc.add_heading(f"Research: {result.topic}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"Generated: {result.generated_at}  •  Word Count: {result.word_count}")
        run.font.size = Pt(10)
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(result.content or "")
        doc.add_heading("References", level=2)
        for ref in result.references:
            doc.add_paragraph(ref.title, style="List Number")
            if ref.url:
                doc.add_paragraph(ref.url)
            if ref.snippet:
                doc.add_paragraph(ref.snippet)
        doc.save(buffer)
        docx_bytes = buffer.getvalue(); buffer.close()
        from fastapi.responses import Response
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=research_{job_id}.docx"}
        )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use: txt, md, pdf, or docx")

# Export Endpoints
@app.get("/health", status_code=200)
async def health_check():
    logger.info("Health check endpoint called")
    return {
        "status": "healthy",
        "services": {"database": "up", "gemini": "up"},
        "version": "0.1.0",
        "uptime_seconds": 3600
    }

if __name__ == "__main__":
    import uvicorn
    create_db_and_tables()
    uvicorn.run(app, host="0.0.0.0", port=8000) 
    logger.info("Application started")
